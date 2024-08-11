"""
DocxProcess 模块

这个模块提供了处理 Word 文档（.docx）的功能，主要用于提取和处理教育培养方案中的课程信息。
主要功能包括：
1. 导入 Word 文档并提取表格和段落信息
2. 提取课程学分信息
3. 处理和格式化表格数据
4. 导出数据到 JSON 文件
5. 从 JSON 文件导入数据

该模块使用 PyQt6 提供图形界面交互，使用 python-docx 处理 Word 文档。
"""

import json
import os
import sys
from collections import namedtuple
from io import BytesIO
import re
from typing import Union

import pandas as pd
from PyQt6.QtWidgets import QFileDialog, QMessageBox, QDialog, QVBoxLayout, QLabel, QPushButton
from docx import Document

sys.path.append(os.getcwd())


class DocxProcess:
    """
    DocxProcess 类用于处理 Word 文档，提取课程信息，并提供导入导出功能。

    属性:
        parent: 父窗口对象
        required_column: 必需的列名
        config_path: 配置文件路径
        last_file_path: 上次使用的文件路径
    """

    def __init__(self, parent=None):
        """
        初始化 DocxProcess 对象。

        :param parent: 父窗口对象，默认为 None
        """
        self.parent = parent
        self.required_column = "理论教学学时"
        self.config_path = os.path.join("..", "config", "user_config.json")
        self.last_file_path = self.load_last_file_path()

    def load_last_file_path(self):
        """
        从配置文件中加载上次使用的文件路径。

        :return: 上次使用的文件路径，如果不存在则返回空字符串
        """
        if os.path.exists(self.config_path):
            with open(self.config_path, 'r') as f:
                config = json.load(f)
                return config.get('education_program_file_path', '')
        return ''

    def save_last_file_path(self, file_path):
        """
        保存最后使用的文件路径到配置文件。

        :param file_path: 要保存的文件路径
        """
        os.makedirs(os.path.dirname(self.config_path), exist_ok=True)

        existing_data = {}
        if os.path.exists(self.config_path):
            try:
                with open(self.config_path, 'r') as f:
                    existing_data = json.load(f)
            except json.JSONDecodeError:
                pass

        existing_data['education_program_file_path'] = file_path

        with open(self.config_path, 'w') as f:
            json.dump(existing_data, f, indent=4, ensure_ascii=False)

    def import_docx(self, student_id):
        """
        导入 Word 文档并处理。
        :param student_id: 导入的文档对应的学号
        :return: 处理后的表格和段落信息，如果导入失败则返回 None
        """
        if self.last_file_path and os.path.exists(self.last_file_path):
            reply = QMessageBox.question(self.parent, '使用上次文件',
                                         f"是否使用上次导入的文件?\n{self.last_file_path}",
                                         QMessageBox.StandardButton.Yes |
                                         QMessageBox.StandardButton.No,
                                         QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                return self.process_file(self.last_file_path, student_id=student_id)

        file_name, _ = QFileDialog.getOpenFileName(
            self.parent,
            "导入 Word 文档",
            "",
            "Word 文档 (*.docx)"
        )

        if not file_name:
            QMessageBox.information(self.parent, "提示", "导入已取消")
            return None

        if not file_name.endswith('.docx'):
            QMessageBox.warning(self.parent, "警告", "请选择 .docx 格式的文件！")
            return None

        self.save_last_file_path(file_name)
        return self.process_file(file_name, student_id=student_id)

    import os

    def process_file(self, file_name, student_id):
        """
        处理 Word 文档文件。

        :param file_name: 要处理的文件路径
        :param student_id: 学生ID
        :return: 处理后的表格和段落信息，如果处理失败则返回 None
        """
        # 获取当前脚本的目录
        script_dir = os.path.dirname(os.path.abspath(__file__))

        # 构建相对于脚本的 JSON 文件路径
        json_file_path = os.path.join(script_dir, "..", "data", f"{student_id}.json")

        # 将路径标准化，解析任何 '..' 和 '.'
        json_file_path = os.path.normpath(json_file_path)

        if not os.path.exists(json_file_path):
            QMessageBox.warning(self.parent, "警告", "未找到成绩数据，请在初始界面进行导入")
            return None

        try:
            with open(file_name, 'rb') as file:
                docx_content = BytesIO(file.read())

            document = Document(docx_content)
            tables_with_paragraphs = self.extract_tables_and_paragraphs(document=document,
                                                                        json_file_path=json_file_path)
            self.export_to_json(results=tables_with_paragraphs, student_id=student_id)

            QMessageBox.information(self.parent, "成功", f"成功导入文件: {file_name}")

            return tables_with_paragraphs, json_file_path

        except Exception as e:
            QMessageBox.critical(self.parent, "错误", f"导入文件时发生错误: {str(e)}")
            return None

    def extract_credit_info(self, strings):
        """
        从字符串中提取课程学分信息。

        :param strings: 包含课程学分信息的字符串列表
        :return: 包含课程类型、必修学分和选修学分的命名元组列表
        """
        CourseInfo = namedtuple('CourseInfo', ['course_type', 'required_credits', 'elective_credits'])

        credit_info = []

        pattern = r'(.*?)\s*最低必修学分数[:：]\s*(\d+)\s*最低选修学分数[:：]\s*(\d+)'

        seen_course_types = set()

        for string in strings:
            match = re.search(pattern, string)
            if match:
                course_type = match.group(1).strip()
                required_credits = int(match.group(2))
                elective_credits = int(match.group(3))

                if course_type not in seen_course_types:
                    credit_info.append(CourseInfo(course_type, required_credits, elective_credits))
                    seen_course_types.add(course_type)

        return credit_info

    def extract_tables_and_paragraphs(self, document, json_file_path):
        """
        从 Word 文档中提取表格和段落信息。
        读取成绩信息，和培养方案合并。
        体育、大学英语课程默认已修读

        :param document: Word 文档对象
        :param json_file_path: JSON 文件路径
        :return: 包含表格数据和相关信息的列表
        """
        results = []
        paragraphs = list(document.paragraphs)
        tables = list(document.tables)

        desired_columns = ['课程名称', '修读形式', '学分', '总学时', '开课学年', '开课学期']

        text = [p.text for p in paragraphs]
        relevant_paragraph = [content for content in text if
                              any(keyword in content for keyword in ["最低选修学分数", "最低必修学分数"])]
        score_need = self.extract_credit_info(relevant_paragraph)

        # 检查 JSON 文件是否存在
        if not os.path.exists(json_file_path):
            dialog = MessageDialog("教务成绩数据不存在，请在主界面进行导入")
            dialog.exec()
            return results

        # 读取 JSON 文件
        try:
            with open(json_file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
        except json.JSONDecodeError:
            dialog = MessageDialog("无法解析教务成绩数据，请在主界面重新导入")
            dialog.exec()
            return results
        except Exception as e:
            dialog = MessageDialog(f"读取教务成绩数据时发生错误：{str(e)}")
            dialog.exec()
            return results

        # 提取课程信息
        course_info = {course['课程名']: course for course in json_data[1:]}

        delete_index = []
        for i, table in enumerate(tables):
            if not table.rows:
                delete_index.append(i)
                continue

            header_row = [cell.text.strip() for cell in table.rows[0].cells]
            if self.required_column not in header_row:
                delete_index.append(i)

        tables = [table for i, table in enumerate(tables) if i not in delete_index]

        for i, table in enumerate(tables):
            original_header_row = [cell.text.strip() for cell in table.rows[0].cells]

            # 获取所需列的索引
            column_indices = {col: original_header_row.index(col) for col in desired_columns if
                              col in original_header_row}

            if '课程名称' not in column_indices:
                dialog = MessageDialog(f"警告: 表格 {i + 1} 中没有找到 '课程名称' 列")
                dialog.exec()
                continue

            # 只保留所需的列，并添加新列
            header_row = [col for col in desired_columns if col in column_indices] + ["状态", "成绩", "绩点"]
            table_data = []

            for row in table.rows[1:]:
                row_data = [row.cells[column_indices[col]].text.strip() for col in desired_columns if
                            col in column_indices]
                course_name = row.cells[column_indices['课程名称']].text.strip()

                if course_name in course_info and course_name != "小计":
                    # 课程已修读
                    status = "已修读"
                    score = course_info[course_name].get('总成绩', '')
                    grade_point = course_info[course_name].get('绩点', '')
                elif course_name in ["体育", "大学英语", "跨学科基本课程", "形势与政策",
                                     "新时代中国特色社会主义劳动教育"]:
                    status = "已修读"
                    score = ''
                    grade_point = ''
                elif course_name == "小计":
                    status = ''
                    score = ''
                    grade_point = ''
                else:
                    # 课程未修读或正在修读（这里简单处理为未修读）
                    status = "未修读"
                    score = ''
                    grade_point = ''

                row_data.extend([status, score, grade_point])
                table_data.append(row_data)

            results.append({
                'table': {
                    'header': header_row,
                    'data': table_data
                },
                'info': score_need[i] if i < len(score_need) else None
            })

        return results

    def export_to_json(self, results, student_id):
        """
        将结果导出到 JSON 文件。

        :param results: 要导出的结果数据
        :param student_id: 成绩数据对应的学号
        """
        json_filename = "degree_progress_" + str(student_id) + ".json"
        config_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'config'))
        os.makedirs(config_dir, exist_ok=True)

        file_path = os.path.join(config_dir, json_filename)

        serializable_results = []
        for item in results:
            serializable_item = {
                'table': item['table'],
                'info': item['info']
            }
            serializable_results.append(serializable_item)

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(serializable_results, f, ensure_ascii=False, indent=4)

        print(f"Results saved to {file_path}")

    def import_from_json(self, json_filename="degree_progress.json"):
        """
        从 JSON 文件导入数据。

        :param json_filename: JSON 文件名，默认为 "degree_progress.json"
        :return: 导入的数据，如果文件不存在则返回 None
        """
        config_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'config'))
        file_path = os.path.join(config_dir, json_filename)

        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None

        with open(file_path, 'r', encoding='utf-8') as f:
            loaded_results = json.load(f)

        results = []
        for item in loaded_results:
            result = {
                'table': item['table'],
                'info': item['info']
            }
            results.append(result)

        print(f"Results loaded from {file_path}")
        return results


class MessageDialog(QDialog):
    def __init__(self, message):
        super().__init__()
        self.setWindowTitle("提示")
        layout = QVBoxLayout()
        label = QLabel(message)
        button = QPushButton("确定")
        button.clicked.connect(self.accept)
        layout.addWidget(label)
        layout.addWidget(button)
        self.setLayout(layout)
